"""Pytest fixtures."""

import pytest
from docx import Document


@pytest.fixture
def sample_docx_path(tmp_path):
    """Create a minimal DOCX with paragraphs and a table."""
    doc = Document()
    doc.add_paragraph("hello - привет")
    doc.add_paragraph("world – мир")
    table = doc.add_table(rows=2, cols=1)
    table.cell(0, 0).text = "apple - яблоко"
    table.cell(1, 0).text = "banana – банан"
    path = tmp_path / "sample.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def docx_with_table(tmp_path):
    """DOCX with a table containing word pairs."""
    doc = Document()
    table = doc.add_table(rows=3, cols=1)
    table.cell(0, 0).text = "word - перевод"
    table.cell(1, 0).text = "phrase – фраза"
    table.cell(2, 0).text = "item — элемент"
    path = tmp_path / "table.docx"
    doc.save(str(path))
    return path

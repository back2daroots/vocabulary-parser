"""Extract text from DOCX files (paragraphs and tables)."""

from pathlib import Path

from docx import Document


def extract_text_from_docx(file_path: Path) -> str:
    """
    Extract text from a DOCX file.
    Iterates paragraphs and all table cells.
    """
    doc = Document(str(file_path))
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)

    return "\n".join(parts) if parts else ""
